"""
parameter:1.plant  工厂用于连接数据库
          2.pmm01  采购单号，用户SQL取数据
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import cx_Oracle as oracle
import sys
import os

"""
#获取当前工作目录 
l_ydir = os.getcwd()

#更换至ftp_word目录
os.chdir("E:/PythonData/Word")
"""
os.environ['NLS_LANG'] = 'AMERICAN_AMERICA.AL32UTF8'
#连接oracle
l_conn_sql = '%(username)s/%(password)s@topprod'\
             %{'username':sys.argv[1],'password':sys.argv[1]}
conn = oracle.connect(l_conn_sql)

#获取单号及供应商数据
cr = conn.cursor()
l_sql="select pmm01,pmm04,pmm09,pmc03,pmm22,pmm43 from pmm_file left join pmc_file on pmc01 = pmm09 where " \
      "pmm01 = \'%s\'"%(sys.argv[2])
print(l_sql)
cr.execute(l_sql)
rs01 = cr.fetchall()
#获取单号的明细数据
l_sql="select pmn04,ima02,ima021,pmn20,pmn07,pmn33,pmn31,pmn31t from pmn_file left join ima_file on pmn04 = ima01 " \
      "where pmn01 = \'%s\'"%(sys.argv[2])
cr.execute(l_sql)
rs02 = cr.fetchall()
cr.close()
conn.close()

#生成word文档
document = Document()

h=document.add_heading('xxx技术股份有限公司',0)
h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph('采购单号:')
p.add_run(rs01[0][0]).bold = True
p.add_run('             下单日期:'+str(rs01[0][1])[0:10])

p1 = document.add_paragraph('供应商:'+rs01[0][3])
p1.add_run('          币种:'+rs01[0][4])
p1.add_run('          税率:'+str(rs01[0][5])+'%')

table = document.add_table(rows=1, cols=8)        #创建表格
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '物料编码'
hdr_cells[1].text = '品名'
hdr_cells[2].text = '规格'
hdr_cells[3].text = '数量'
hdr_cells[4].text = '单位'
hdr_cells[5].text = '交货日期'
hdr_cells[6].text = '未税单价'
hdr_cells[7].text = '含税单价'
for pmn04, ima02, ima021,pmn20,pmn07,pmn33,pmn31,pmn31t in rs02:
    row_cells = table.add_row().cells
    row_cells[0].text = pmn04
    row_cells[1].text = ima02
    row_cells[2].text = ima021
    row_cells[3].text = str(pmn20)
    row_cells[4].text = pmn07
    row_cells[5].text = str(pmn33)[0:10]
    row_cells[6].text = str(pmn31)
    row_cells[7].text = str(pmn31t)

document.add_page_break()

#按照‘单号_供应商.word’格式保存
print(type(rs01[0][0]),rs01[0][2])
document.save("E:/PythonData/%s_%s.docx"%(rs01[0][0],rs01[0][3]))  #保存文档