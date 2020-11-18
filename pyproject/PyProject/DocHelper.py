from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT   #用作设置段落对齐
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
from docx.shared import Inches #图片尺寸
import time

document =Document()
today = time.strftime("%Y-%m-%d", time.localtime())

def InitReportDocument():
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal'].font.size = Pt(13)
    # 设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    #添加图片   图片路径和尺寸
    # document.add_picture('1.jpg',width=Inches(6))
def AddDocumentTitle(docTitle):
    p1 = document.add_paragraph()
    # 初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
    p1.aligment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 标题内容
    run1 = p1.add_run(docTitle)
    run1.font.name = '微软雅黑'
    run1.font.size = Pt(21)  # 设置字体
    # 设置加粗
    run1.font.bold = True
    # 段后距离5磅
    p1.space_after = Pt(5)
    ##段前距离5磅
    p1.space_before = Pt(5)

def AddDocumentDesc(reportDesc):
    p2 = document.add_paragraph()
    run2 = p2.add_run(reportDesc)  # 添加总述
    run2.font.name = '微软雅黑'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(13)  # 设置字体
    # 设置加粗
    run2.font.bold = True

def AddTableDesc(faultDesc, size):
    p3 = document.add_paragraph()
    run3 = p3.add_run(faultDesc)  # 表格总说明或描述内容
    run3.font.name = '仿宋_GB2312'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(size)  # 设置字体

def CreateTable(title, cellList):
    p4 = document.add_paragraph()
    # 添加一个表格  行列 和格式
    table = document.add_table(rows=2, cols=len(cellList), style='Table Grid')
    # 合并单元格
    table.cell(0, 0).merge(table.cell(0, len(cellList) - 1))
    # 对于合并的单元格，输入其中任何一个单元格都可以
    table_run1 = table.cell(0, 0).paragraphs[0].add_run(title)  # 表格标题
    table_run1.font.name = u'微软雅黑'
    table_run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    table.cell(0, 0).paragraphs[0].aligment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 使用默认字体和格式
    iCell = 0
    for colName in cellList:
        table.cell(1, iCell).text = colName
        iCell = iCell + 1
    # table.cell(1, 1).text = colName2
    # table.cell(1, 2).text = colName3
    # table.cell(1, 3).text = colName4

def AddTableRow(cellList):
    table = document.add_table(rows=1, cols=len(cellList), style='Table Grid')
    iCell = 0
    for value in cellList:
        table.cell(0, iCell).text = value
        iCell = iCell + 1
    # table.cell(0, 0).text = colValue1
    # table.cell(0, 1).text = colValue2
    # table.cell(0, 2).text = colValue3
    # table.cell(0, 3).text = colValue4
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(12)

def SaveReport(outputFile):
    document.save(outputFile)
#插入分页符
# document.add_page_break()