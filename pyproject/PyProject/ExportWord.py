from docx import Document
from docx.shared import Inches, Pt

'''
模板内容
【测试标题】如PVE故障注入自动测试
【测试记录】测试项目，测试日期，测试人员
【测试故障注入】
故障名称、故障描述、注入类型、故障注入结果，如添加电阻 电阻挡位 普通电阻 注入成功
【INCA操作】
 操作类型（读取/标定） 变量名称 设定值  当前值
 【OBDII】
 MODE类型 MODE查询结果
'''

doc = Document()  # 以默认模板建立文档对象
# doc = Document('a.docx')  # 读取a.docx文档，建立文档对象

def chg_font(obj, fontname='微软雅黑', size=None):
    ## 设置字体函数
    obj.font.name = fontname
    # obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    if size and isinstance(size, Pt):
        obj.font.size = size

distance = Inches(3)
sec = doc.sections[0]  # sections对应文档中的“节”
sec.left_margin = distance  # 以下依次设置左、右、上、下页面边距
sec.right_margin = distance
sec.top_margin = distance
sec.bottom_margin = distance
sec.page_width = Inches(12)  # 设置页面宽度
sec.page_height = Inches(20)  # 设置页面高度
##设置默认字体
# chg_font(doc.styles['Normal'], fontname='宋体')

doc.add_heading('Document Title', level=1)
# 添加段落文本
paragraph = doc.add_paragraph('this is paragraph')
peir_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
paragraph = doc.add_paragraph('add content')
paragraph.add_run('doclor sit amet')
ph_format = paragraph.paragraph_format
ph_format.space_before = Pt(10)  # 设置段前间距
ph_format.space_after = Pt(12)  # 设置段后间距
ph_format.line_spacing = Pt(19)  # 设置行间距

# 设置同一个段落的不同文本格式
run = paragraph.add_run('text...')
run.bold = True  # 设置字体为粗体
chg_font(run, fontname='微软雅黑', size=Pt(12))  # 设置字体和字号

#添加表格
# tab = doc.add_table(rows=4, cols=4)  # 添加一个4行4列的空表
# cell = tab.cell(1, 3)  # 获取某单元格对象（从0开始索引）
# #单元格添加文本
# cell.text = 'abc'

# 单元格添加多行文本
# ph = cell.paragraphs[0]
# run = ph.add_run('text....')
# run.add_break()  # 添加一个折行
# run.add_picture('a.png')  # 插入图像，可以是内存中的图像，width=Inches(1.0)指定宽度。

#文档种添加图像
# doc.add_picture('a.png')

# 保存
doc.save('a.docx')